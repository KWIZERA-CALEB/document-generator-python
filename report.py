from docx import Document
from docx.shared import RGBColor, Inches, Pt

document = Document()

document.add_heading('Report Cards')

paragraph1 = document.add_paragraph('Hello world')

image1 = document.add_picture('pic.png', width=Inches(3))

table = document.add_table(rows=3,cols=4, style="Table Grid")

first_row = table.rows[0].cells
first_row[0].text = "First Name"
first_row[1].text = "Last Name"
first_row[2].text = "Age"
first_row[3].text = "Email"

second_row = table.rows[1].cells
second_row[0].text = "Kwizera"
second_row[1].text = "Caleb"
second_row[2].text = "20"
second_row[3].text = "kwizeracaleb91@gmail.com"

third_row = table.rows[2].cells
third_row[0].text = "Seth"
third_row[1].text = "Abijuru"
third_row[2].text = "30"
third_row[3].text = "sethabi@gmail.com"


paragraph2 = document.add_paragraph('Hi bro')
run = paragraph2.runs[0]
run.font.color.rgb = RGBColor(255, 0, 0)

run.bold = True
# paragraph2.font.size = Pt(30)

document.save('demo.docx')


