from docx import Document
import warnings
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
import random

# import openai
import openai


warnings.filterwarnings('ignore', category=UserWarning, module='docx.styles.styles')

document = Document()

# suggested comments turples
bad_comments = [
    "It's disappointing you didn't pass, but don't be discouraged. Let's work on improvement together.",
    "I know this isn't the result you hoped for. Let's figure out how to turn this around.",
    "Falling short is tough, but with perseverance and a plan, you can succeed next time.",
    "You didn't pass this time, but failure is a stepping stone to success. Let's learn from this and move forward.",
    "This setback doesn't define you. Let's identify the areas that need improvement and work on them.",
    "I understand you're disappointed. Let's discuss what went wrong and how we can improve.",
    "Not passing is tough, but remember it's a temporary setback. Let's work on a strategy for improvement.",
    "It's okay to feel disappointed. Let's use this as motivation to come back stronger next time.",
    "Failing is part of the learning process. Let's focus on how you can bounce back from this.",
    "I know you're capable of more. Let's work together to turn this failure into a success story."
]

medium_comments = [
    "Good effort, you're halfway there! Keep pushing.",
    "You're making progress. Continue to work hard and you'll get there.",
    "Nice try! With a bit more effort, you can achieve even better results.",
    "You're on the right track. Keep up the good work and aim higher.",
    "Your hard work is evident. Stay focused and you can improve even more.",
    "Well done for reaching halfway. Keep striving for improvement.",
    "You're showing potential. With continued effort, you'll reach your goals.",
    "Halfway there! Keep studying and practicing to improve further.",
    "You've made a good start. Keep up the effort and you'll see better results.",
    "Good job on your progress so far. Continue to work hard and you will succeed."
]

good_comments = [
    "Great job on passing! Your hard work has paid off.",
    "Congratulations on passing! Keep up the excellent work.",
    "Well done! You've successfully passed. Continue to aim high.",
    "You did it! Your dedication has led to your success.",
    "Fantastic work! You've passed with flying colors.",
    "Excellent effort! Passing this subject is a great achievement.",
    "Congratulations on your success! Keep building on this momentum.",
    "You passed! Your commitment and hard work are evident.",
    "Great effort! Passing is just the beginning of your journey.",
    "Well done on passing! Your perseverance has been rewarded."
]
# suggested comments turples


# Adding Date paragraph
date_paragraph = document.add_paragraph('2023/2024, Term 2')
date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
date_run = date_paragraph.runs[0]
date_run.font.size = Pt(16)
date_run.bold = True
date_run.font.name = 'Calibri'

# add image
student_avatar = document.add_picture('pic.png', width=Inches(1))

# Adding Students Names paragraph
students_names_paragraph = document.add_paragraph()
students_names_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
students_names_run = students_names_paragraph.add_run('NAMES:')
students_names_run.font.size = Pt(16)
students_names_run.bold = True
students_names_run.font.name = 'Calibri'

# Adding individual student names
learners_name = input('Learners Name:')
learners_name.upper()
student_name_run = students_names_paragraph.add_run(f'{learners_name}')
student_name_run.font.size = Pt(16)
student_name_run.font.underline = True
student_name_run.bold = True
student_name_run.font.color.rgb = RGBColor(10, 10, 200)
student_name_run.font.name = 'Calibri'

# Adding Grade paragraph
learners_grade = input('Grade:')
student_grade_paragraph = document.add_paragraph(f'GRADE: {learners_grade}')
student_grade_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
student_grade_run = student_grade_paragraph.runs[0]
student_grade_run.font.size = Pt(16)
student_grade_run.bold = True
student_grade_run.font.name = 'Calibri'

# page break
# document.add_page_break()


teachers_comment_title = document.add_paragraph('Home room teacher’s comment')
teachers_comment_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
teachers_comment_title_run = teachers_comment_title.runs[0]
teachers_comment_title_run.font.size = Pt(16)
teachers_comment_title_run.bold = True
teachers_comment_title_run.font.name = 'Calibri'



# create table
table1 = document.add_table(rows=4, cols=4, style="Table Grid")

# fill it 
first_row = table1.rows[0].cells
first_row[0].text = 'Activities'
first_row[1].text = 'Usually'
first_row[2].text = 'Sometimes'
first_row[3].text = 'Rarely'

second_row = table1.rows[1].cells
second_row[0].text = 'Well-prepared and punctual'
second_row[1].text = '✓'
second_row[2].text = 'x'
second_row[3].text = '✓'

third_row = table1.rows[2].cells
third_row[0].text = 'Completes assignments, meets deadlines'
third_row[1].text = '✓'
third_row[2].text = 'x'
third_row[3].text = '✓'

fourth_row = table1.rows[3].cells
fourth_row[0].text = 'Follows class and school rules'
fourth_row[1].text = '✓'
fourth_row[2].text = 'x'
fourth_row[3].text = '✓'


# increase the width for column 1
for row in table1.rows:
    row.cells[0].width = Inches(6)




# table two
table2 = document.add_table(rows=1, cols=4)
row_one = table2.rows[0].cells
row_one[0].text = 'Days: 6'
row_one[1].text = 'Absent: 2'
row_one[2].text = 'Late: 2'
row_one[3].text = 'P.E.Non-Suit: NULL'




comment_teachers_room = document.add_paragraph(f'Comment:  {learners_name} has made great progress in her learning, now reading short sentences independently. She enjoys school activities with friends and am proud of her.Excited for next term.')
comment_teachers_room.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
comment_teachers_room_run = comment_teachers_room.runs[0]
comment_teachers_room_run.font.size = Pt(12)
comment_teachers_room_run.font.name = 'Bookman Old'




teachers_name = document.add_paragraph('Teacher Namara')
teachers_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
teachers_name_run = teachers_name.runs[0]
teachers_name_run.font.size = Pt(12)
teachers_name_run.font.name = 'Bookman Old'

school_name = document.add_paragraph('Two wings international school ')
school_name.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
school_name_run = school_name.runs[0]
school_name_run.font.size = Pt(14)
school_name_run.bold = True
school_name_run.font.name = 'Calibri'

exam_details_title = document.add_paragraph('End of term 3 report: 2023-2024 Term 2')
exam_details_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
exam_details_title_run = exam_details_title.runs[0]
exam_details_title_run.font.size = Pt(12)
exam_details_title_run.font.name = 'Bookman Old'

# student_details = document.add_paragraph()
# student_details.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
# student_details_run = student_details.runs[1]
# student_details_run.font.size = Pt(12)
# student_details_run.font.name = 'Bookman Old'
# student_details_add_names_title = student_details.add_run('Student:')
# student_details_add_names = student_details.add_run('Anika Teta')


the_subject = input('The Subject:')
library_title = document.add_paragraph(f'{the_subject}')
library_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
library_title_run = library_title.runs[0]
library_title_run.font.size = Pt(14)
library_title_run.bold = True
library_title_run.font.name = 'Calibri'

# table3 = document.add_table(rows=2, cols=2)

# classwork_title = document.add_paragraph('Classwork')
# classwork_title_run = classwork_title.runs[0]

# custom variables
classwork = input('Classwork-First-Subject:')
total_classwork = 30

homework = input('Homework-First-Subject:')
total_homework = 20

exams = input('Exams-First-Subject:')
total_exams = 50

sum = 100
total = (int(classwork) + int(homework) + int(exams)) / sum

total_average = total * 100


if total_average == 50:
    comment_from_teacher = random.choice(medium_comments)
    learners_effort = 5
    learners_attitude = 6
elif total_average < 50:
    comment_from_teacher = random.choice(bad_comments)
    learners_effort = 4
    learners_attitude = 3
elif total_average > 50:
    comment_from_teacher = random.choice(good_comments)
    learners_effort = 10
    learners_attitude = 9



table4 = document.add_table(rows=2, cols=3)
table4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
table4_row_one = table4.rows[0].cells
table4_row_one[0].text = 'Classwork'
table4_row_one[1].text = 'Homework'
table4_row_one[2].text = 'Exams'

table4_row_two = table4.rows[1].cells
table4_row_two[0].text = f'{classwork}/30'
table4_row_two[1].text = f'{homework}/20'
table4_row_two[2].text = f'{exams}/50'

total_average_title = document.add_paragraph()
total_average_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
total_average_title_run = total_average_title.add_run('Total:')
total_average_title_run.font.size = Pt(12)
total_average_title_run.bold = True
total_average_title_run.font.name = 'Calibri'

# add total marks
total_average_title_run = total_average_title.add_run(f'{int(total_average)}%')
total_average_title_run.font.size = Pt(12)
total_average_title_run.bold = True
total_average_title_run.font.name = 'Calibri'

table5 = document.add_table(rows=1, cols=3)
table5_row_one = table5.rows[0].cells
table5_row_one[0].text = f'Grade: {learners_grade}'
table5_row_one[1].text = f'Effort: {learners_effort}'
table5_row_one[2].text = f'Attitude: {learners_attitude}'



# create an instance of Openai
# client = openai.OpenAI(
#     # api_key=''
#     api_key = ''
# )

# {"role": "user", "content": f'Create a good comment for this student with {classwork} out of {total_classwork} in classwork, {homework} out of {total_homework}, {exams} out of {total_exams} and also has the average of {total_average}'}

# use openai
# completion = client.chat.completions.create(
#     model = 'gpt-3.5-turbo',
#     messages = [
#         {"role": "user", "content": "Hello"}
#     ],
#     max_tokens=10
# )

# comment_from_openai = completion.choices[0].message['content']



subject_comment = document.add_paragraph(f'Comment: {learners_name} {comment_from_teacher}')





course_content_title = document.add_paragraph('Course Outline')
course_content_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
course_content_title_run = course_content_title.runs[0]
course_content_title_run.font.size = Pt(14)
course_content_title_run.bold = True
course_content_title_run.font.name = 'Calibri'

course_contents_input = input('Course contents(Separate by comma):')


# split the course contents by comma

course_contents = course_contents_input.split(', ')


for course_content in course_contents:
    document.add_paragraph(f'{course_content}', style="ListBullet")

teacher_izina = input('Teacher Name:')
subject_teacher_name = document.add_paragraph(f'Tr.{teacher_izina}')
subject_teacher_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER





# for second subject
second_subject = input('Second Subject:')

sec_sub_title = document.add_paragraph(f'{second_subject}')
sec_sub_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
sec_sub_title_run = sec_sub_title.runs[0]
sec_sub_title_run.font.size = Pt(14)
sec_sub_title_run.bold = True
sec_sub_title_run.font.name = 'Calibri'

classwork_second_subject = input('Classwork-Second-Suject:')
total_classwork_second_subject = 30
homework_second_subject = input('Homework-Second-Suject:')
total_homework_second_subject = 20
exam_second_subject = input('Exam-Second-Suject:')
total_exam_second_subject = 50

sum_second_subject = 100

total_second_subject = (int(classwork_second_subject) + int(homework_second_subject) + int(exam_second_subject)) / sum_second_subject

total_average_second_subject = total_second_subject * 100

if total_average_second_subject == 50:
    comment_from_teacher_sec_sub =  random.choice(medium_comments)
    learners_effort_sec_sub = 5
    learners_attitude_sec_sub = 6
elif total_average_second_subject < 50:
    comment_from_teacher_sec_sub = random.choice(bad_comments)
    learners_effort_sec_sub = 4
    learners_attitude_sec_sub = 3
elif total_average_second_subject > 50:
    comment_from_teacher_sec_sub = random.choice(good_comments)
    learners_effort_sec_sub = 10
    learners_attitude_sec_sub = 9

table6 = document.add_table(rows=2, cols=3)
table6.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
table6_row_one = table6.rows[0].cells
table6_row_one[0].text = 'Classwork'
table6_row_one[1].text = 'Homework'
table6_row_one[2].text = 'Exams'

table6_row_two = table6.rows[1].cells
table6_row_two[0].text = f'{classwork_second_subject}/30'
table6_row_two[1].text = f'{homework_second_subject}/20'
table6_row_two[2].text = f'{exam_second_subject}/50'

total_average_sec_sub_title = document.add_paragraph()
total_average_sec_sub_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
total_average_sec_sub_title_run = total_average_sec_sub_title.add_run('Total:')
total_average_sec_sub_title_run.font.size = Pt(12)
total_average_sec_sub_title_run.bold = True
total_average_sec_sub_title_run.font.name = 'Calibri'

# add total second subject marks
total_average_sec_sub_title_run = total_average_sec_sub_title.add_run(f'{int(total_average_second_subject)}%')
total_average_sec_sub_title_run.font.size = Pt(12)
total_average_sec_sub_title_run.bold = True
total_average_sec_sub_title_run.font.name = 'Calibri'

table7 = document.add_table(rows=1, cols=3)
table7_row_one = table7.rows[0].cells
table7_row_one[0].text = f'Grade: {learners_grade}'
table7_row_one[1].text = f'Effort: {learners_effort_sec_sub}'
table7_row_one[2].text = f'Attitude: {learners_attitude_sec_sub}'

subject_sec_sub_comment = document.add_paragraph(f'Comment: {learners_name} {comment_from_teacher_sec_sub}')

# course contents for second subject
course_two_content_title = document.add_paragraph('Course Outline')
course_two_content_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
course_two_content_title_run = course_two_content_title.runs[0]
course_two_content_title_run.font.size = Pt(14)
course_two_content_title_run.bold = True
course_two_content_title_run.font.name = 'Calibri'

course_two_contents_input = input('Course2 contents(Separate by comma):')


# split the course contents by comma

course_two_contents = course_two_contents_input.split(', ')


for course_two_content in course_two_contents:
    document.add_paragraph(f'{course_two_content}', style="ListBullet")





teacher_izina_sec_sub = input('Second Teacher Name:')
subject_teacher_name_sec_sub = document.add_paragraph(f'Tr.{teacher_izina_sec_sub}')
subject_teacher_name_sec_sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER




# total sum
total_sum_title = document.add_paragraph('Total Marks')
total_sum_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
total_sum_title_run = total_sum_title.runs[0]
total_sum_title_run.font.size = Pt(14)
total_sum_title_run.bold = True
total_sum_title_run.underline = True
total_sum_title_run.font.name = 'Calibri'

total_sum_add = total + total_second_subject
total_sum_all = 2

total_sum_all_subs_add_divide = total_sum_add / total_sum_all
total_sum_all_avg = total_sum_all_subs_add_divide * 100

all_together_total = document.add_paragraph(f'Total Sum: {total_sum_all_avg}%')
all_together_total.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
all_together_total_run = all_together_total.runs[0]
all_together_total_run.font.size = Pt(14)
all_together_total_run.font.name = 'Calibri'
all_together_total_run.bold = True

if total_sum_all_avg < 50:
    final_given_comment = f'{learners_name} Try to improve next term'
elif total_sum_all_avg == 50:
    final_given_comment = f'{learners_name} Put in much effort you can make it'
elif total_sum_all_avg > 50:
    final_given_comment = f'{learners_name} Keep it up'


final_comment = document.add_paragraph(f'{final_given_comment}')


# Save the document
document.save(f'{learners_name}-Report.docx')


# import csv

# # Initialize empty lists for each type of comments
# encouraging_comments = []
# passing_comments = []
# failing_comments = []

# Read the CSV file
# with open('comments.csv', mode='r') as file:
#     csv_reader = csv.DictReader(file)
#     for row in csv_reader:
#         comment_type = row['type']
#         comment = row['comment']
        
#         # Append the comment to the respective list based on its type
#         if comment_type == 'encouraging':
#             encouraging_comments.append(comment)
#         elif comment_type == 'passing':
#             passing_comments.append(comment)
#         elif comment_type == 'failing':
#             failing_comments.append(comment)